import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import io
import hashlib
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import reporte_cartera
import ficha_caso
import acta_inspeccion
from tablero_datos import parsear_fecha_flexible

# --- SECCIÓN 1: CONFIGURACIÓN Y MOTOR DE CARGA ---
st.set_page_config(page_title="Dashboard de Gestión: Procesos y Tendencias", layout="wide", page_icon="⚙️")

@st.cache_data(ttl=300, show_spinner="Descargando Pipeline desde Google Sheets...")
def cargar_pipeline_google_cacheado():
    return reporte_cartera.cargar_pipeline_desde_google_sheets(dict(st.secrets["gcp_service_account"]))

@st.cache_data
def generar_datos_prueba():
    np.random.seed(42)
    hoy = datetime(2026, 4, 1) 
    areas = ["Ingeniería y Energía", "Equipos Móviles"]
    liquidadores = ["Carlos Mendoza", "Ana Rojas", "Luis Silva", "Marta Pérez"]
    estados_principales = ["Ingreso", "Instrucción y Análisis", "Resolución", "Cerrado"]
    
    fechas_ingreso = [hoy - timedelta(days=int(x)) for x in np.random.randint(2, 1000, 1500)]
    datos = []
    
    for f in fechas_ingreso:
        area = np.random.choice(areas, p=[0.3, 0.7]) 
        liquidador = np.random.choice(liquidadores)
        dias_desde_ingreso = (hoy - f).days
        es_cerrado = np.random.random() < 0.8
        estado = "Cerrado" if es_cerrado else "En Análisis"
        
        datos.append({
            "ID_Caso": f"CASO-{np.random.randint(100000, 999999)}",
            "Area_Negocio": area,
            "Liquidador": liquidador,
            "Estado_Actual": estado,
            "Subestado_Actual": "Subestado de prueba",
            "Fecha_Ingreso": f,
            "Fecha_Cierre": f + timedelta(days=np.random.randint(5, 60)) if es_cerrado else pd.NaT,
            "Días desde asignación": np.random.randint(1, 100),
            "Días desde contacto": np.random.randint(1, 10),
            "Días entre inspección e asignación": np.random.randint(1, 15),
            "Días informe inicial - inspección": np.random.randint(1, 20),
            "Días análisis contractual - asignación": np.random.randint(5, 45),
            "Días informe final despachado - análisis contractual enviado": np.random.randint(1, 30),
            "Perdida bruta (en moneda del caso)": np.random.randint(1000, 50000),
            "Monto asegurado (en moneda del caso)": np.random.randint(50000, 1000000),
            "Gastos (UF)": np.random.uniform(5, 50),
            "Honorarios (UF)": np.random.uniform(10, 100)
        })
    return pd.DataFrame(datos)

def buscar_indice_columna(columnas, palabras_clave):
    for i, col in enumerate(columnas):
        if str(col).strip().lower() in palabras_clave:
            return i
    for i, col in enumerate(columnas):
        for palabra in palabras_clave:
            if palabra in str(col).strip().lower():
                return i
    return 0

@st.cache_data(show_spinner="Extrayendo fotos del Acta de Inspección...")
def extraer_fotos_acta_cacheado(archivo_bytes, nombre_archivo):
    return acta_inspeccion.extraer_fotos_acta(archivo_bytes, nombre_archivo)

@st.cache_data(show_spinner="Buscando Hechos y Circunstancias en el Acta...")
def extraer_descripcion_siniestro_cacheado(archivo_bytes, nombre_archivo):
    return acta_inspeccion.extraer_descripcion_siniestro(archivo_bytes, nombre_archivo)

def _google_configurado():
    """`"x" in st.secrets` lanza StreamlitSecretNotFoundError si no existe
    ningún secrets.toml (p.ej. antes de configurar los secretos en un deploy
    nuevo); esto evita que tumbe toda la app."""
    try:
        return "gcp_service_account" in st.secrets
    except Exception:
        return False

@st.cache_data(ttl=300, show_spinner="Descargando Base Maestra desde Google Sheets...")
def cargar_base_maestra_google_cacheado():
    df = reporte_cartera.cargar_hoja_desde_google_sheets(
        dict(st.secrets["gcp_service_account"]), "Base_Maestra", header_row=2
    )
    return df.rename(columns={
        "Número de caso": "ID_Caso",
        "División": "Area_Negocio",
        "Ajustador senior": "Liquidador",
        "Estado": "Estado_Actual",
        "Sub estado": "Subestado_Actual",
        "Creado en": "Fecha_Ingreso",
        "Fecha de cierre": "Fecha_Cierre",
    })

st.sidebar.title("Configuración y Carga")

df_raw = None
conectado_google_maestra = _google_configurado()

if conectado_google_maestra:
    col_bm1, col_bm2 = st.sidebar.columns([3, 1])
    with col_bm1:
        st.caption("🔗 Base Maestra conectada desde Google Sheets.")
    with col_bm2:
        if st.button("🔄", key="btn_actualizar_base_maestra", help="Actualizar Base Maestra"):
            cargar_base_maestra_google_cacheado.clear()
    try:
        df_raw = cargar_base_maestra_google_cacheado()
    except Exception as e:
        st.sidebar.error(f"No se pudo leer la Base Maestra desde Google Sheets: {e}")

archivo_subido = st.sidebar.file_uploader(
    "Cargar Reporte de Casos (CSV/Excel) — respaldo manual" if conectado_google_maestra else "Cargar Reporte de Casos (CSV/Excel)",
    type=["csv", "xlsx"],
)

if df_raw is not None:
    pass
elif archivo_subido is not None:
    try:
        filas_saltar = st.sidebar.number_input("Filas a saltar (Encabezado desfasado)", min_value=0, max_value=20, value=5)
        
        if archivo_subido.name.endswith('.csv'):
            df_crudo = pd.read_csv(archivo_subido, skiprows=filas_saltar, low_memory=False)
        else:
            xl = pd.ExcelFile(archivo_subido)
            hoja_seleccionada = st.sidebar.selectbox("Selecciona la pestaña de tu Excel", xl.sheet_names)
            df_crudo = pd.read_excel(archivo_subido, sheet_name=hoja_seleccionada, skiprows=filas_saltar)
            
        columnas_reales = df_crudo.columns.tolist()
        st.sidebar.success("¡Archivo detectado! Mapeo automático activado.")
        
        idx_id = buscar_indice_columna(columnas_reales, ['número de caso', 'numero de caso', 'id'])
        idx_area = buscar_indice_columna(columnas_reales, ['división', 'division', 'área de negocio'])
        idx_liq = buscar_indice_columna(columnas_reales, ['ajustador senior', 'liquidador'])
        idx_estado = buscar_indice_columna(columnas_reales, ['estado'])
        idx_subestado = buscar_indice_columna(columnas_reales, ['sub estado', 'subestado'])
        idx_in = buscar_indice_columna(columnas_reales, ['creado en', 'fecha de denuncio'])
        idx_out = buscar_indice_columna(columnas_reales, ['fecha de cierre', 'fecha cierre'])

        col_id = st.sidebar.selectbox("Columna ID Caso", columnas_reales, index=idx_id)
        col_area = st.sidebar.selectbox("Columna División", columnas_reales, index=idx_area)
        col_liq = st.sidebar.selectbox("Columna Ajustador Senior", columnas_reales, index=idx_liq)
        col_estado = st.sidebar.selectbox("Columna Estado", columnas_reales, index=idx_estado)
        col_subestado = st.sidebar.selectbox("Columna Sub estado", columnas_reales, index=idx_subestado)
        col_fecha_in = st.sidebar.selectbox("Columna Creado en", columnas_reales, index=idx_in)
        col_fecha_out = st.sidebar.selectbox("Columna Fecha Cierre", columnas_reales, index=idx_out)

        # Renombramos solo lo fundamental para el sistema, el resto pasa tal cual
        df_raw = df_crudo.rename(columns={
            col_id: "ID_Caso",
            col_area: "Area_Negocio",
            col_liq: "Liquidador",
            col_estado: "Estado_Actual",
            col_subestado: "Subestado_Actual",
            col_fecha_in: "Fecha_Ingreso",
            col_fecha_out: "Fecha_Cierre"
        })
        
    except Exception as e:
        st.sidebar.error(f"Error al procesar el archivo: {e}")
        df_raw = generar_datos_prueba()
else:
    st.sidebar.info("Usando datos de demostración para visualización.")
    df_raw = generar_datos_prueba()

# --- SECCIÓN 2: MOTOR DE CÁLCULO ESTRICTO ---
def procesar_datos_integrales(df):
    # BLINDAJE ANTI-KEYERROR: Asegura que las columnas existan aunque el Excel falle
    for col in ['Estado_Actual', 'Subestado_Actual', 'Area_Negocio', 'Liquidador']:
        if col not in df.columns:
            df[col] = 'Desconocido'
    for col in ['Fecha_Ingreso', 'Fecha_Cierre']:
        if col not in df.columns:
            df[col] = pd.NaT

    df['Estado_Actual'] = df['Estado_Actual'].fillna('Desconocido').astype(str).str.strip().str.upper()
    df['Subestado_Actual'] = df['Subestado_Actual'].fillna('Desconocido').astype(str).str.strip().str.upper()
    df['Area_Negocio'] = df['Area_Negocio'].fillna('Sin Área').astype(str).str.strip()
    
    # 1. FILTRO: Omitir casos Rechazados
    df = df[~df['Estado_Actual'].str.contains('RECHAZADO|RECHAZO', case=False, na=False)]
    df = df[~df['Subestado_Actual'].str.contains('RECHAZADO|RECHAZO', case=False, na=False)]
    
    # 2. DETECCIÓN Y FILTRO DE TIEMPOS DE RESIDENCIA (Columnas "Días")
    # Si el encabezado quedó desfasado, alguna columna puede llegar como fecha/número en vez de texto.
    cols_dias = [col for col in df.columns if isinstance(col, str) and ('Días' in col or 'Dias' in col)]
    for c in cols_dias:
        df[c] = pd.to_numeric(df[c], errors='coerce').fillna(0)
        # FILTRO: Eliminar errores lógicos (outliers de > 1500 días)
        df = df[df[c] < 1500]

    # Transformación de Fechas Clave (acepta texto normal o número de serie de
    # Excel, como llega 'Fecha de cierre' en la Base Maestra de Google Sheets)
    df['Fecha_Ingreso'] = df['Fecha_Ingreso'].apply(parsear_fecha_flexible)
    df['Fecha_Cierre'] = df['Fecha_Cierre'].apply(parsear_fecha_flexible)
    
    df['Es_Abierto'] = ~df['Estado_Actual'].str.contains('CERRADO') & df['Fecha_Cierre'].isna()
    
    df['Mes_Cierre'] = df['Fecha_Cierre'].dt.to_period('M').astype(str)
    df['Trimestre_Cierre'] = df['Fecha_Cierre'].dt.to_period('Q').astype(str)
    df['Año_Cierre'] = df['Fecha_Cierre'].dt.year.astype(str).replace('nan', 'Pendiente')
    
    df_abiertos = df[df['Es_Abierto']].copy()
    df_cerrados = df[~df['Es_Abierto']].copy()
            
    return df_abiertos, df_cerrados, df, cols_dias

df_abiertos, df_cerrados, df_master, columnas_de_dias = procesar_datos_integrales(df_raw)

# --- FILTROS DE INTERFAZ TEMPORAL ---
st.sidebar.header("Filtros de Tendencias")
tipo_periodo = st.sidebar.radio("Agrupación Temporal:", ["Mensual", "Trimestral", "Anual"])

if tipo_periodo == "Mensual":
    periodos_disp = sorted(df_cerrados['Mes_Cierre'].unique(), reverse=True) if not df_cerrados.empty else []
    col_cierre = 'Mes_Cierre'
elif tipo_periodo == "Trimestral":
    periodos_disp = sorted(df_cerrados['Trimestre_Cierre'].unique(), reverse=True) if not df_cerrados.empty else []
    col_cierre = 'Trimestre_Cierre'
else:
    periodos_disp = sorted(df_cerrados['Año_Cierre'].unique(), reverse=True) if not df_cerrados.empty else []
    col_cierre = 'Año_Cierre'

periodos_limpios = [p for p in periodos_disp if p != 'NaT' and p != 'Pendiente' and p != 'nan']
periodo_seleccionado = st.sidebar.selectbox("Seleccionar Periodo Final:", periodos_limpios) if periodos_limpios else None

# --- SECCIÓN 3: MOTOR DE REPORTE VISUAL (DASHBOARD) ---
st.title("📊 Panel de Gestión: Tiempos de Residencia Reales")
st.markdown("Monitor de control basado exclusivamente en los registros de tiempo de tu sistema.")

# --- Carga del Pipeline (compartida entre "Reporte de Cartera" y "Por Ajustador") ---
st.markdown("#### 🔗 Pipeline (para los reportes PPTX)")

df_pipeline = None
conectado_google = _google_configurado()

col_fuente, col_actualizar = st.columns([4, 1])
with col_fuente:
    if conectado_google:
        st.caption("🔗 Pipeline conectado automáticamente desde Google Sheets.")
    else:
        st.caption("Conexión a Google Sheets no configurada todavía; sube el archivo manualmente.")
with col_actualizar:
    if conectado_google and st.button("🔄 Actualizar Pipeline", use_container_width=True, key="btn_actualizar_pipeline"):
        cargar_pipeline_google_cacheado.clear()

if conectado_google:
    try:
        df_pipeline = cargar_pipeline_google_cacheado()
    except Exception as e:
        st.error(
            "No se pudo leer el Pipeline desde Google Sheets. Verifica que la hoja esté "
            f"compartida como Lector con la cuenta de servicio. Detalle: {e}"
        )

if df_pipeline is None:
    archivo_pipeline = st.file_uploader(
        "Cargar Pipeline (Excel) — respaldo manual" if conectado_google else "Cargar Pipeline (Excel)",
        type=["xlsx"], key="archivo_pipeline",
    )
    if archivo_pipeline is not None:
        df_pipeline = reporte_cartera.cargar_pipeline(archivo_pipeline)

st.divider()


def renderizar_editor_y_pptx(df_filtrado, titulo_sugerido, titulo_key, key_prefix, titulo_label):
    """Editor de la tabla de casos + generación del pptx. Compartido entre el
    Reporte de Cartera (filtrado por Corredora/Aseguradora/Asegurado) y el
    Reporte por Ajustador (filtrado por Ajustador senior): misma plantilla y
    lógica, solo cambia cómo se filtró el Pipeline antes de llegar aquí."""
    if df_filtrado.empty:
        st.info("No hay casos que cumplan con los filtros seleccionados.")
        return

    col_t1, col_t2 = st.columns([2, 1])
    with col_t1:
        titulo_cartera = st.text_input(titulo_label, value=titulo_sugerido, key=titulo_key)
    with col_t2:
        fecha_corte_reporte = st.date_input("Fecha de corte del reporte", value=datetime.now().date(), key=f"{key_prefix}_fecha_corte")

    tabla_base = reporte_cartera.preparar_tabla_casos(df_filtrado, fecha_corte_reporte)

    st.markdown(f"**Cartera filtrada: {len(tabla_base)} casos activos.** "
                "La probabilidad de cierre y la observación ya vienen del Pipeline; ajústalas si es necesario antes de generar el pptx.")

    tabla_editada = st.data_editor(
        tabla_base[["Caso", "Numero_Siniestro", "Asegurado", "Nickname", "Divisa", "Perdida_bruta", "Dias", "Division", "MCL", "Prob", "Observacion", "Observacion_sugerida"]],
        column_config={
            "Perdida_bruta": st.column_config.NumberColumn("Pérdida bruta", format="%.0f", disabled=True),
            "Dias": st.column_config.NumberColumn("Días", disabled=True),
            "Division": st.column_config.TextColumn("División", disabled=True),
            "MCL": st.column_config.CheckboxColumn("MCL", disabled=True),
            "Caso": st.column_config.TextColumn("Caso", disabled=True),
            "Numero_Siniestro": st.column_config.TextColumn("N° Siniestro", disabled=True),
            "Asegurado": st.column_config.TextColumn("Asegurado", disabled=True),
            "Divisa": st.column_config.TextColumn("Divisa", disabled=True),
            "Prob": st.column_config.NumberColumn("Prob. cierre (%)", min_value=0, max_value=100, step=25),
            "Observacion_sugerida": st.column_config.TextColumn("Observación original (referencia)", disabled=True),
        },
        num_rows="fixed",
        use_container_width=True,
        key=f"{key_prefix}_tabla_editada",
        height=350,
    )

    st.markdown("#### Próximos pasos y focos de gestión (hasta 5)")
    pasos = []
    for i in range(5):
        c1, c2 = st.columns([1, 2])
        with c1:
            titulo = st.text_input(f"Título paso {i + 1}", key=f"{key_prefix}_paso_titulo_{i}")
        with c2:
            desc = st.text_input(f"Descripción paso {i + 1}", key=f"{key_prefix}_paso_desc_{i}")
        if titulo.strip() or desc.strip():
            pasos.append({"titulo": titulo, "desc": desc})

    alerta_prioritaria = st.text_area(
        "Alerta de atención prioritaria (slide 'Casos que requieren atención especial')",
        key=f"{key_prefix}_alerta",
        placeholder="Ej: Atención prioritaria: Caso XXXXX (Nickname · Prob. 0%) acumula USD > X M sin antecedentes.",
    )

    if st.button("🎯 Generar PPTX", use_container_width=True, key=f"{key_prefix}_btn_generar"):
        pptx_bytes = reporte_cartera.generar_pptx(
            fecha_corte_reporte,
            titulo_cartera,
            pd.DataFrame(tabla_editada),
            pasos,
            alerta_prioritaria,
        )
        nombre_archivo = "".join(c if c.isalnum() else "_" for c in titulo_cartera).strip("_") or "Cartera"
        st.download_button(
            label="⬇️ Descargar Estado_Cartera.pptx",
            data=pptx_bytes,
            file_name=f"Estado_Cartera_{nombre_archivo}_{fecha_corte_reporte.strftime('%d%m%y')}.pptx",
            mime="application/vnd.openxmlformats-officedocument.presentationml.presentation",
            use_container_width=True,
            key=f"{key_prefix}_btn_descargar",
        )


tab_reporte, tab_ajustador, tab_energia, tab_moviles, tab_tendencias, tab_ficha = st.tabs([
    "📑 Reporte Ejecutivo de Cartera (PPTX)", "👤 Reporte por Ajustador (PPTX)",
    "⚡ WIP: Ingeniería y Energía", "🚜 WIP: Equipos Móviles", "📈 Cierres e Históricos",
    "🔎 Ficha de Caso (PPTX)",
])

with tab_reporte:
    st.subheader("📑 Reporte Ejecutivo de Cartera (PPTX)")

    if df_pipeline is None:
        st.info("Conecta el Pipeline (arriba) desde Google Sheets o sube el archivo para generar el reporte.")
    else:

        st.markdown("#### Panorama general del Pipeline (Top 5)")
        col_g1, col_g2, col_g3 = st.columns(3)
        top5_specs = [
            (col_g1, "Compañía de seguros", "Top 5 Aseguradoras", "#3498db"),
            (col_g2, "Corredora", "Top 5 Corredoras", "#e67e22"),
            (col_g3, "Asegurado", "Top 5 Asegurados", "#27ae60"),
        ]
        for columna, campo, titulo_grafico, color in top5_specs:
            with columna:
                st.markdown(f"**{titulo_grafico}**")
                if campo in df_pipeline.columns:
                    top5 = df_pipeline[campo].dropna().value_counts().head(5).reset_index()
                    top5.columns = [campo, "Casos"]
                    fig_top5 = px.bar(
                        top5.sort_values("Casos"), x="Casos", y=campo, orientation="h",
                        text="Casos", color_discrete_sequence=[color],
                    )
                    fig_top5.update_layout(yaxis_title=None, height=280, margin=dict(l=0, r=10, t=10, b=10))
                    st.plotly_chart(fig_top5, use_container_width=True)
                else:
                    st.info(f"El Pipeline no tiene la columna '{campo}'.")

        st.divider()

        col_f1, col_f2, col_f3 = st.columns(3)
        with col_f1:
            opciones_corredora = sorted(df_pipeline["Corredora"].dropna().unique().tolist()) if "Corredora" in df_pipeline.columns else []
            filtro_corredoras = st.multiselect("Corredor", opciones_corredora, key="filtro_corredoras")
        with col_f2:
            opciones_aseguradora = sorted(df_pipeline["Compañía de seguros"].dropna().unique().tolist()) if "Compañía de seguros" in df_pipeline.columns else []
            filtro_aseguradoras = st.multiselect("Aseguradora", opciones_aseguradora, key="filtro_aseguradoras")
        with col_f3:
            opciones_asegurado = sorted(df_pipeline["Asegurado"].dropna().unique().tolist()) if "Asegurado" in df_pipeline.columns else []
            filtro_asegurados = st.multiselect("Asegurado", opciones_asegurado, key="filtro_asegurados")

        if not (opciones_corredora or opciones_aseguradora or opciones_asegurado):
            st.warning("El archivo cargado no tiene columnas 'Corredora' / 'Compañía de seguros' / 'Asegurado'; se usará todo el Pipeline sin filtrar.")

        df_cartera_filtrada = reporte_cartera.filtrar_pipeline(
            df_pipeline, corredoras=filtro_corredoras, aseguradoras=filtro_aseguradoras, asegurados=filtro_asegurados
        )

        titulo_sugerido = reporte_cartera.sugerir_titulo_cartera(filtro_corredoras, filtro_aseguradoras, filtro_asegurados)
        # La key incluye los filtros activos para que el sugerido se refresque cada vez que cambian,
        # sin perder una edición manual mientras el usuario no toque los filtros.
        titulo_key = "titulo_cartera__" + "|".join(sorted(filtro_corredoras + filtro_aseguradoras + filtro_asegurados))
        renderizar_editor_y_pptx(
            df_cartera_filtrada, titulo_sugerido, titulo_key, key_prefix="cartera",
            titulo_label="Título de la cartera (portada, resumen y pie de página)",
        )

with tab_ajustador:
    st.subheader("👤 Reporte de Cartera por Ajustador (PPTX)")
    st.caption("Mismo reporte que 'Reporte Ejecutivo de Cartera', filtrado por Ajustador Senior — para ver el detalle de la carga de trabajo de cada ajustador.")

    if df_pipeline is None:
        st.info("Conecta el Pipeline (arriba) desde Google Sheets o sube el archivo para generar el reporte.")
    elif "Ajustador senior" not in df_pipeline.columns:
        st.warning("El Pipeline no tiene la columna 'Ajustador senior'.")
    else:
        carga_ajustador = df_pipeline["Ajustador senior"].dropna().value_counts().reset_index()
        carga_ajustador.columns = ["Ajustador senior", "Casos"]

        st.markdown("#### Carga de casos activos por Ajustador Senior")
        fig_carga_aj = px.bar(
            carga_ajustador.sort_values("Casos"), x="Casos", y="Ajustador senior", orientation="h",
            text="Casos", color_discrete_sequence=["#8e44ad"],
        )
        fig_carga_aj.update_layout(yaxis_title=None, height=max(280, 28 * len(carga_ajustador)), margin=dict(l=0, r=10, t=10, b=10))
        st.plotly_chart(fig_carga_aj, use_container_width=True)

        st.divider()

        opciones_ajustador = sorted(df_pipeline["Ajustador senior"].dropna().unique().tolist())
        filtro_ajustadores = st.multiselect("Ajustador Senior", opciones_ajustador, key="filtro_ajustadores")

        if not filtro_ajustadores:
            st.info("Selecciona uno o más Ajustadores para generar su reporte de carga de trabajo.")
        else:
            df_cartera_ajustador = reporte_cartera.filtrar_pipeline(df_pipeline, ajustadores=filtro_ajustadores)
            titulo_sugerido_aj = reporte_cartera.sugerir_titulo_ajustador(filtro_ajustadores)
            titulo_key_aj = "titulo_ajustador__" + "|".join(sorted(filtro_ajustadores))
            renderizar_editor_y_pptx(
                df_cartera_ajustador, titulo_sugerido_aj, titulo_key_aj, key_prefix="ajustador",
                titulo_label="Título del reporte (portada, resumen y pie de página)",
            )

def renderizar_panel_area(df_area_abiertos, area_nombre, cols_dias):
    if df_area_abiertos.empty:
        st.success(f"No hay casos activos detectados para la división {area_nombre}.")
        return
    
    st.markdown(f"**Total Casos en Curso (WIP): {len(df_area_abiertos)}**")
    
    col1, col2 = st.columns([1, 1])
    with col1:
        st.markdown(f"#### Carga por Ajustador Senior ({area_nombre})")
        carga_liq = df_area_abiertos.groupby('Liquidador').size().reset_index(name='Casos Asignados')
        carga_liq = carga_liq.sort_values('Casos Asignados', ascending=True)
        fig_carga = px.bar(carga_liq, x='Casos Asignados', y='Liquidador', orientation='h', text='Casos Asignados', color_discrete_sequence=['#3498db'])
        st.plotly_chart(fig_carga, use_container_width=True)

    with col2:
        st.markdown("#### Tiempos Promedio de Residencia (Casos Activos)")
        st.markdown("Promedio de días registrados por el sistema en las etapas de vida del caso.")
        
        # Filtramos las columnas que realmente tienen datos para graficar los promedios
        promedios_dias = []
        for c in cols_dias:
            promedio = df_area_abiertos[c].mean()
            if promedio > 0:
                promedios_dias.append({'Etapa (Columna del Sistema)': c, 'Días Promedio': promedio})
                
        if promedios_dias:
            df_promedios = pd.DataFrame(promedios_dias).sort_values('Días Promedio', ascending=True)
            fig_dias = px.bar(df_promedios, x='Días Promedio', y='Etapa (Columna del Sistema)', orientation='h', text_auto='.1f', color_discrete_sequence=['#e74c3c'])
            st.plotly_chart(fig_dias, use_container_width=True)
        else:
            st.info("No hay datos de 'Días' registrados para los casos activos de esta división.")

with tab_energia:
    renderizar_panel_area(df_abiertos[df_abiertos['Area_Negocio'].str.contains('Ingeniería', case=False, na=False)], 'Ingeniería y Energía', columnas_de_dias)

with tab_moviles:
    renderizar_panel_area(df_abiertos[df_abiertos['Area_Negocio'].str.contains('Móvil|Movil', case=False, na=False)], 'Equipos Móviles', columnas_de_dias)

with tab_tendencias:
    st.subheader(f"Análisis Retrospectivo ({tipo_periodo})")
    if not df_cerrados.empty and periodo_seleccionado:
        df_cierre_periodo = df_cerrados[df_cerrados[col_cierre] == periodo_seleccionado]
        
        c1, c2 = st.columns(2)
        with c1:
            st.metric("Total Resoluciones del Periodo", len(df_cierre_periodo))
            vol_energia = len(df_cierre_periodo[df_cierre_periodo['Area_Negocio'].str.contains('Ingeniería', case=False, na=False)])
            st.markdown(f"**Ingeniería:** {vol_energia} casos")
        with c2:
            st.metric("Total Ajustadores Involucrados", df_cierre_periodo['Liquidador'].nunique())
            vol_moviles = len(df_cierre_periodo[df_cierre_periodo['Area_Negocio'].str.contains('Móvil|Movil', case=False, na=False)])
            st.markdown(f"**Móviles:** {vol_moviles} casos")
            
        st.divider()
        
        st.markdown("#### Tendencias Históricas de Cierre (Volumen)")
        historico_agrupado = df_cerrados.groupby([col_cierre, 'Area_Negocio']).agg(Volumen=('ID_Caso', 'count')).reset_index().sort_values(col_cierre)
        ultimos_periodos = sorted([p for p in historico_agrupado[col_cierre].unique() if p != 'NaT' and p != 'nan'])[-12:]
        historico_filtrado = historico_agrupado[historico_agrupado[col_cierre].isin(ultimos_periodos)]

        fig_vol = px.line(historico_filtrado, x=col_cierre, y='Volumen', color='Area_Negocio', markers=True)
        st.plotly_chart(fig_vol, use_container_width=True)
    else:
        st.info("No hay casos cerrados con fechas válidas para mostrar tendencias históricas.")

def buscar_observacion_caso(caso_id, df_pipeline, fila_base_maestra):
    """Busca la observación completa para 'Estado Actual del Siniestro':
    primero en el Pipeline (por Caso JPV), y si no hay match ahí, cae a la
    Base Maestra del propio caso. Devuelve (texto, fuente) o (None, None)."""
    if df_pipeline is not None and "Número de caso" in df_pipeline.columns:
        match = df_pipeline[df_pipeline["Número de caso"].astype(str) == str(caso_id)]
        if not match.empty:
            fila_pipeline = match.iloc[0]
            texto = reporte_cartera.limpiar_observacion(fila_pipeline.get("Observaciones", ""))
            if not texto or texto.lower() in ("nan", "none"):
                texto = reporte_cartera.limpiar_observacion(fila_pipeline.get("Contenido último movimiento", ""))
            if texto and texto.lower() not in ("nan", "none"):
                return texto, "Observaciones del Pipeline"

    for campo in ["Observaciones", "Última observación", "Contenido último movimiento"]:
        texto = reporte_cartera.limpiar_observacion(fila_base_maestra.get(campo, ""))
        if texto and texto.lower() not in ("nan", "none"):
            return texto, "Base Maestra"
    return None, None


with tab_ficha:
    st.subheader("🔎 Ficha de Caso Individual")
    st.caption("Busca un caso por N° de Caso JPV o por texto libre en el Nickname para generar su ficha (pptx de 3 slides: resumen, registro fotográfico y gestiones, estas dos últimas para completar a mano).")

    if df_master.empty or "ID_Caso" not in df_master.columns:
        st.info("No hay datos cargados para buscar un caso.")
    else:
        col_b1, col_b2 = st.columns(2)
        with col_b1:
            opciones_caso = sorted(df_master["ID_Caso"].dropna().astype(str).unique().tolist())
            caso_directo = st.selectbox("Caso JPV", ["(todos)"] + opciones_caso, key="ficha_caso_directo")
        with col_b2:
            busqueda_nickname = st.text_input("Buscar por Nickname (texto libre)", key="ficha_busqueda_nickname")

        candidatos = df_master.copy()
        candidatos["ID_Caso"] = candidatos["ID_Caso"].astype(str)
        if caso_directo != "(todos)":
            candidatos = candidatos[candidatos["ID_Caso"] == caso_directo]
        if busqueda_nickname.strip() and "Nickname" in candidatos.columns:
            candidatos = candidatos[candidatos["Nickname"].astype(str).str.contains(busqueda_nickname.strip(), case=False, na=False)]

        if candidatos.empty:
            st.info("No se encontraron casos con esos criterios.")
        else:
            if len(candidatos) > 1:
                nicknames = candidatos["Nickname"].astype(str) if "Nickname" in candidatos.columns else ""
                etiquetas = (candidatos["ID_Caso"] + " — " + nicknames).tolist()
                elegido = st.selectbox(f"Se encontraron {len(candidatos)} casos, elige uno:", etiquetas, key="ficha_desambiguar")
                fila = candidatos.iloc[etiquetas.index(elegido)]
            else:
                fila = candidatos.iloc[0]

            st.divider()
            st.markdown(f"### {fila.get('Nickname') or fila.get('Asegurado') or fila.get('ID_Caso')}")
            divisa_fila = str(fila.get("Divisa") or "").strip()
            col_r1, col_r2 = st.columns(2)
            with col_r1:
                st.markdown(f"**Asegurado:** {fila.get('Asegurado', '—')}")
                st.markdown(f"**N° Siniestro:** {fila.get('Número de siniestro', '—')}")
                st.markdown(f"**Caso JPV:** {fila.get('ID_Caso', '—')}")
                st.markdown(f"**Estado:** {fila.get('Estado_Actual', '—')}")
                st.markdown(f"**Monto asegurado:** {fila.get('Monto asegurado (en moneda del caso)', '—')} {divisa_fila}")
            with col_r2:
                st.markdown(f"**Pérdida bruta:** {fila.get('Perdida bruta (en moneda del caso)', '—')} {divisa_fila}")
                st.markdown(f"**Fecha de ocurrencia:** {fila.get('Fecha de ocurrencia', '—')}")
                st.markdown(f"**Fecha de denuncio:** {fila.get('Fecha de denuncio', '—')}")
                st.markdown(f"**Fecha de asignación:** {fila.get('Fecha de asignación', '—')}")
                st.markdown(f"**Días desde asignación:** {fila.get('Días desde asignación', '—')}")

            st.divider()
            st.markdown("#### 📷 Acta de Inspección (opcional)")
            st.caption(
                "Sube el Acta de Inspección (Word o PDF) para precargar el Registro Fotográfico y la "
                "Descripción del Siniestro (desde \"Hechos y Circunstancias\"). Word es más confiable; "
                "PDF funciona solo si el acta fue exportada desde Word (no un PDF escaneado) y es "
                "mejor esfuerzo — revisa el resultado antes de dar por buena la ficha."
            )
            archivo_acta = st.file_uploader("Cargar Acta de Inspección (.docx o .pdf)", type=["docx", "pdf"], key="ficha_acta_uploader")

            fotos_seleccionadas = []
            descripcion_siniestro_texto, descripcion_siniestro_fuente = None, None
            if archivo_acta is not None:
                acta_bytes = archivo_acta.getvalue()
                acta_id = hashlib.md5(acta_bytes).hexdigest()[:8]
                try:
                    fotos_candidatas = extraer_fotos_acta_cacheado(acta_bytes, archivo_acta.name)
                except Exception as e:
                    fotos_candidatas = []
                    st.error(f"No se pudo leer el Acta de Inspección: {e}")

                try:
                    descripcion_siniestro_texto = extraer_descripcion_siniestro_cacheado(acta_bytes, archivo_acta.name)
                    if descripcion_siniestro_texto:
                        descripcion_siniestro_fuente = "Hechos y Circunstancias del Acta de Inspección"
                except Exception:
                    descripcion_siniestro_texto = None

                if descripcion_siniestro_texto:
                    st.success("Se encontró \"Hechos y Circunstancias\" en el Acta — se precargará en Descripción del Siniestro.")

                if not fotos_candidatas:
                    st.warning("No se encontraron fotos en el Acta cargada.")
                else:
                    st.markdown(f"**{len(fotos_candidatas)} fotos encontradas.** Desmarca las que no sirvan para la ficha.")
                    cols_galeria = st.columns(3)
                    for i, foto in enumerate(fotos_candidatas):
                        with cols_galeria[i % 3]:
                            st.image(foto["imagen"], use_container_width=True)
                            marcado = st.checkbox(
                                foto["pie"][:70] or f"Foto {i + 1}", value=True, key=f"ficha_foto_check_{acta_id}_{i}",
                            )
                            if marcado:
                                fotos_seleccionadas.append(foto)

            if st.button("🎯 Generar Ficha de Caso (PPTX)", use_container_width=True, key="btn_generar_ficha"):
                observacion_texto, observacion_fuente = buscar_observacion_caso(fila.get("ID_Caso"), df_pipeline, fila)
                pptx_bytes = ficha_caso.generar_ficha_pptx(
                    fila, observacion_texto, observacion_fuente, fotos_seleccionadas or None,
                    descripcion_siniestro_texto, descripcion_siniestro_fuente,
                )
                nombre_archivo = "".join(c if c.isalnum() else "_" for c in str(fila.get("ID_Caso", ""))).strip("_") or "Caso"
                st.download_button(
                    label="⬇️ Descargar Ficha_Caso.pptx",
                    data=pptx_bytes,
                    file_name=f"Ficha_Caso_{nombre_archivo}.pptx",
                    mime="application/vnd.openxmlformats-officedocument.presentationml.presentation",
                    use_container_width=True,
                    key="btn_descargar_ficha",
                )

# --- SECCIÓN 4: MOTOR DE REPORTES EXPORTABLES (EXCEL Y WORD) ---
st.divider()
st.subheader("📥 Generación de Reportes Formales")

@st.cache_data
def generar_excel_completo(df_master, df_abiertos):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_master.to_excel(writer, index=False, sheet_name='Base_Filtrada_Limpia')
        df_abiertos.to_excel(writer, index=False, sheet_name='WIP_Abiertos')
    return output.getvalue()

def generar_grafico_mpl(df, x_col, y_col, titulo, ylabel, color):
    plt.figure(figsize=(7, 3.5))
    plt.plot(df[x_col], df[y_col], marker='o', color=color, linewidth=2)
    plt.title(titulo, fontsize=10, fontweight='bold')
    plt.ylabel(ylabel, fontsize=9)
    plt.xticks(rotation=45, ha='right', fontsize=8)
    plt.grid(axis='y', linestyle='--', alpha=0.7)
    plt.tight_layout()
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=120)
    plt.close()
    img_stream.seek(0)
    return img_stream

def generar_word_reporte(df_abiertos, df_cerrados, periodo_sel, col_cierre):
    doc = Document()
    doc.add_heading('Reporte Ejecutivo de Operaciones e Ingeniería', 0)
    doc.add_paragraph(f'Periodo de Análisis y Corte: {periodo_sel}')
    
    doc.add_heading('1. Estado Actual del Portafolio (WIP)', level=1)
    doc.add_paragraph(f'Total de casos en curso a la fecha: {len(df_abiertos)} casos.')
    doc.add_paragraph('Nota: Este reporte excluye todos los casos clasificados como "Rechazados" y anomalías del sistema para asegurar la integridad de las métricas.')
    
    doc.add_heading(f'2. Cierres y Tendencias', level=1)
    if not df_cerrados.empty and periodo_sel:
        datos_periodo = df_cerrados[df_cerrados[col_cierre] == periodo_sel]
        doc.add_paragraph(f'Volumen resuelto en el periodo final: {len(datos_periodo)} casos.')
        
        tendencia = df_cerrados.groupby(col_cierre).size().reset_index(name='Volumen').tail(6)
        if len(tendencia) > 1:
            img_trend = generar_grafico_mpl(tendencia, col_cierre, 'Volumen', 'Evolución de Resoluciones (Últimos periodos)', 'Cantidad de Casos', '#27ae60')
            doc.add_picture(img_trend, width=Inches(6.0))
            
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

col_d1, col_d2 = st.columns(2)

with col_d1:
    excel_data = generar_excel_completo(df_master, df_abiertos)
    st.download_button(label="📊 Descargar Base Limpia (Excel)", data=excel_data,
                       file_name=f"Base_Datos_Limpia_{datetime.now().strftime('%Y%m%d')}.xlsx",
                       mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)

with col_d2:
    if periodo_seleccionado and not df_cerrados.empty:
        word_data = generar_word_reporte(df_abiertos, df_cerrados, periodo_seleccionado, col_cierre)
        st.download_button(label="📝 Generar Reporte de Gerencia (Word)", data=word_data,
                           file_name=f"Reporte_Gerencia_{periodo_seleccionado}.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
    else:
        st.info("No hay información suficiente para generar el reporte Word.")
