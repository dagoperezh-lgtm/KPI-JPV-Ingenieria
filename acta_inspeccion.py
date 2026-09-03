"""
Extrae información desde un Acta de Inspección (.docx o .pdf) para precargar
la Ficha de Caso:

- Fotos + pie de foto -> Registro Fotográfico.
- "Hechos y Circunstancias" -> Descripción del Siniestro.

Word (.docx) es la fuente más confiable: la estructura de tablas separa
limpiamente imagen(es)/texto por celda. PDF (.pdf) es best-effort: se apoya
en la posición (x, y) de cada imagen y bloque de texto en la página, válido
solo si el PDF viene "exportado" desde Word (imágenes y texto como objetos
separados) — un PDF escaneado no tiene nada extraíble de esta forma.

Se expone extraer_fotos_acta()/extraer_descripcion_siniestro() como punto de
entrada único: detectan el formato por la extensión del archivo y usan la
implementación que corresponda.
"""
import io
import os
import re

from docx import Document
from docx.oxml.ns import qn

try:
    import pymupdf
except ImportError:
    pymupdf = None


# ---------------------------------------------------------------------------
# Word (.docx)
# ---------------------------------------------------------------------------

def _imagenes_en_celda(cell):
    rids = []
    for blip in cell._element.findall(".//" + qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if rid:
            rids.append(rid)
    return rids


def _celdas_distintas_con_columnas(table):
    """Recorre la tabla y devuelve, para cada celda DISTINTA (una celda
    combinada horizontalmente aparece repetida en varias posiciones de
    row.cells; acá se reporta una sola vez), en qué columnas de esa fila
    aparece — necesario para saber a qué imagen le corresponde qué pie de
    foto cuando conviven varias en la misma tabla."""
    resultado = []  # [(cols_ocupadas: set[int], cell)]
    for row in table.rows:
        tcs_fila = [c._tc for c in row.cells]  # referencias vivas: comparar por identidad, no por id()
        vistas = []
        for c, cell in enumerate(row.cells):
            if any(cell._tc is tc for tc in vistas):
                continue
            vistas.append(cell._tc)
            cols_ocupadas = {j for j, tc in enumerate(tcs_fila) if tc is cell._tc}
            resultado.append((cols_ocupadas, cell))
    return resultado


def _centro(cols_ocupadas):
    return sum(cols_ocupadas) / len(cols_ocupadas)


def extraer_fotos_acta_docx(archivo):
    """archivo: ruta o file-like (.docx). Devuelve una lista de dicts
    {"imagen": bytes, "pie": str} en el orden en que aparecen en el
    documento.

    Cada imagen se empareja con el pie de foto cuya columna esté más cerca
    (no "el primer texto de la tabla"): varios bloques traen 2+ fotos con
    pies de foto DISTINTOS uno al lado del otro (columnas separadas), y no
    una sola descripción compartida.

    Cualquier tabla cuya primera fila tenga texto (p.ej. "Antecedentes
    Generales", "Firmas") se descarta completa: son tablas de
    secciones/datos o de firmas, no de fotos.
    """
    doc = Document(archivo)
    fotos = []

    for table in doc.tables:
        if len(table.rows) == 0:
            continue
        fila0 = table.rows[0]
        if any(cell.text.strip() for cell in fila0.cells):
            continue

        celdas_imagen = []
        celdas_texto = []
        for cols_ocupadas, cell in _celdas_distintas_con_columnas(table):
            rids = _imagenes_en_celda(cell)
            if rids:
                celdas_imagen.append((cols_ocupadas, rids))
                continue
            texto = cell.text.strip()
            if texto:
                celdas_texto.append((cols_ocupadas, texto))

        for cols_img, rids in celdas_imagen:
            pie = ""
            if celdas_texto:
                centro_img = _centro(cols_img)
                _, pie = min(celdas_texto, key=lambda ct: abs(_centro(ct[0]) - centro_img))
            for rid in rids:
                parte = doc.part.related_parts[rid]
                fotos.append({"imagen": parte.blob, "pie": pie})

    return fotos


_CAMPO_HECHOS_RE = re.compile(r"^hechos\s+y\s+circunstancias\s*:?$", re.IGNORECASE)


def extraer_descripcion_siniestro_docx(archivo):
    """Busca en las tablas del acta la celda "Hechos y Circunstancias:" y
    devuelve el texto completo de la celda vecina, o None si no se
    encuentra."""
    doc = Document(archivo)
    for table in doc.tables:
        for row in table.rows:
            for c, cell in enumerate(row.cells):
                if _CAMPO_HECHOS_RE.match(cell.text.strip()):
                    if c + 1 < len(row.cells):
                        texto = row.cells[c + 1].text.strip()
                        if texto:
                            return texto
    return None


# ---------------------------------------------------------------------------
# PDF (best-effort: solo funciona si el PDF trae imágenes y texto como
# objetos separados, es decir si viene exportado desde Word y no escaneado)
# ---------------------------------------------------------------------------

_PIE_PAGINA_RE = re.compile(r"divisi[oó]n.{0,40}siniestro\s*n[°º]?.{0,40}caso\s*n[°º]?", re.IGNORECASE | re.DOTALL)
_SECCION_FIRMAS_RE = re.compile(r"\bfirmas\b|firma\s+(asegurado|inspector)\s*:", re.IGNORECASE)
_INICIO_HECHOS_RE = re.compile(r"hechos\s+y\s+circunstancias\s*:?", re.IGNORECASE)
_FIN_HECHOS_RE = re.compile(r"lugar\s+del\s+siniestro\s*:", re.IGNORECASE)


def _centro_bbox(bbox):
    return ((bbox[0] + bbox[2]) / 2, (bbox[1] + bbox[3]) / 2)


def extraer_fotos_acta_pdf(archivo):
    """archivo: ruta o file-like (.pdf). Ver extraer_fotos_acta_docx — misma
    idea, pero emparejando por posición (x, y) en la página en vez de
    columna de tabla, ya que un PDF no conserva la estructura de tabla."""
    if pymupdf is None:
        raise RuntimeError("Falta la librería 'pymupdf' (revisa requirements.txt).")
    doc = pymupdf.open(stream=archivo.read(), filetype="pdf") if hasattr(archivo, "read") else pymupdf.open(archivo)
    fotos = []

    for page in doc:
        bloques_todos = [b for b in page.get_text("blocks") if b[4].strip()]
        if any(_SECCION_FIRMAS_RE.search(b[4]) for b in bloques_todos):
            break  # de acá en adelante es la sección de firmas, no hay más fotos del siniestro

        imagenes = sorted(
            (info for info in page.get_image_info(xrefs=True)
             if (info["bbox"][2] - info["bbox"][0]) > 80 and (info["bbox"][3] - info["bbox"][1]) > 80),
            key=lambda info: (round(info["bbox"][1]), info["bbox"][0]),
        )
        if not imagenes:
            continue
        bloques = [b for b in bloques_todos if not _PIE_PAGINA_RE.search(b[4])]

        for info in imagenes:
            cx, cy = _centro_bbox(info["bbox"])
            mejor_pie, mejor_score = "", None
            for b in bloques:
                tx, ty = _centro_bbox(b[:4])
                dy = ty - info["bbox"][3]  # distancia vertical desde el borde inferior de la imagen
                # el pie de foto casi siempre va debajo; una coincidencia arriba se penaliza fuerte
                score = (dy if dy >= -5 else abs(dy) * 4) * 3 + abs(tx - cx)
                if mejor_score is None or score < mejor_score:
                    mejor_score, mejor_pie = score, b[4].strip()
            imagen_bytes = doc.extract_image(info["xref"])["image"]
            fotos.append({"imagen": imagen_bytes, "pie": mejor_pie})

    return fotos


def extraer_descripcion_siniestro_pdf(archivo):
    """Ver extraer_descripcion_siniestro_docx. En PDF el campo puede venir
    partido en varios bloques de texto (incluso cruzando un salto de
    página); se concatenan desde que aparece la etiqueta "Hechos y
    Circunstancias:" hasta el siguiente campo conocido ("Lugar del
    Siniestro:")."""
    if pymupdf is None:
        raise RuntimeError("Falta la librería 'pymupdf' (revisa requirements.txt).")
    doc = pymupdf.open(stream=archivo.read(), filetype="pdf") if hasattr(archivo, "read") else pymupdf.open(archivo)

    capturando = False
    partes = []
    for page in doc:
        for b in page.get_text("blocks"):
            texto = b[4]
            if not texto.strip() or _PIE_PAGINA_RE.search(texto):
                continue
            if not capturando:
                m = _INICIO_HECHOS_RE.search(texto)
                if not m:
                    continue
                cola = texto[m.end():].strip()
                if cola:
                    partes.append(re.sub(r"\s+", " ", cola))
                capturando = True
                continue
            m_fin = _FIN_HECHOS_RE.search(texto)
            if m_fin:
                if m_fin.start() > 0:
                    partes.append(re.sub(r"\s+", " ", texto[: m_fin.start()]).strip())
                return "\n\n".join(p for p in partes if p).strip() or None
            partes.append(re.sub(r"\s+", " ", texto).strip())
    return "\n\n".join(p for p in partes if p).strip() or None


# ---------------------------------------------------------------------------
# Despachador por extensión de archivo
# ---------------------------------------------------------------------------

def _es_pdf(nombre_archivo):
    return os.path.splitext(nombre_archivo or "")[1].lower() == ".pdf"


def extraer_fotos_acta(archivo_bytes, nombre_archivo):
    """archivo_bytes: contenido crudo del archivo (.docx o .pdf).
    nombre_archivo: nombre original, para detectar el formato por extensión.
    Envuelve los bytes en un stream nuevo en cada llamada (así se puede
    invocar junto a extraer_descripcion_siniestro sobre el mismo archivo
    sin preocuparse por la posición de lectura de un stream compartido)."""
    origen = io.BytesIO(archivo_bytes)
    return extraer_fotos_acta_pdf(origen) if _es_pdf(nombre_archivo) else extraer_fotos_acta_docx(origen)


def extraer_descripcion_siniestro(archivo_bytes, nombre_archivo):
    origen = io.BytesIO(archivo_bytes)
    return (
        extraer_descripcion_siniestro_pdf(origen)
        if _es_pdf(nombre_archivo)
        else extraer_descripcion_siniestro_docx(origen)
    )
